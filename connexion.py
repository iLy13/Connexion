from telegram.ext import Updater, MessageHandler, Filters, ConversationHandler
from telegram.ext import CommandHandler
from telegram import ReplyKeyboardMarkup, ReplyKeyboardRemove, message, bot
from functions import start, help, create_kb
from io import BytesIO
from docx import Document
import os
archiv = ''
files = [['/help']]
markup = ReplyKeyboardMarkup(files, one_time_keyboard=False)
chat_id = '1303744874'
n_pages = 0


def get(update, context):
    global archiv, chat_id
    format = ' '.join(context.args).lower().split('.')[-1]
    if format == 'txt':
        with open(f'{archiv}/{" ".join(context.args)}', 'r', encoding='utf-8') as f:
            data = f.read()
            update.message.reply_text(data)
    elif format in ['jpg', 'jpeg', 'png']:
        context.bot.send_photo(chat_id=chat_id, photo=open(f'{archiv}/{" ".join(context.args)}', 'rb'))
    elif format == 'docx':
        document = Document(f'{archiv}/{" ".join(context.args)}')
        for par in document.paragraphs:
            update.message.reply_text(par.text)
        for table in document.tables:
            update.message.reply_text(table.rows[0].cells[0].text)
    else:
        context.bot.send_document(chat_id=chat_id, document=open(f'{archiv}/{" ".join(context.args)}', 'rb'))


def connect(update, context):
    global archiv, files, markup, n_pages
    try:
        n_pages = 0
        archiv = context.args[0]
        files = create_kb(os.listdir(archiv))
        markup = ReplyKeyboardMarkup(files[0], one_time_keyboard=False)
        update.message.reply_text('We are ready to work! Choose file', reply_markup=markup)
    except IndexError:
        update.message.reply_text('You need to write name of the directory')
    except FileNotFoundError:
        update.message.reply_text('You need to write correct name of the directory')


def disconnect(update, context):
    global archiv, files, markup, n_pages
    archiv = ''
    files = [['/help']]
    n_pages = 0
    markup = ReplyKeyboardMarkup(files, one_time_keyboard=False)
    update.message.reply_text("Directory is disconnected. Check our service's commands", reply_markup=markup)


def next_page(update, context):
    global archiv, files, markup, n_pages
    n_pages += 1
    if n_pages == len(files):
        n_pages = 0
    markup = ReplyKeyboardMarkup(files[n_pages], one_time_keyboard=False)
    update.message.reply_text('Next page! Choose file', reply_markup=markup)


def new_text(update, context):
    update.message.reply_text('Enter name of file and in next message print your text')
    return 1


def first_response(update, context):
    data = update.message.text
    print(1)
    context.user_data['file'] = data
    if data.split('.')[-1] == 'txt':
        update.message.reply_text('Print your text')
        return 'txt'
    elif data.split('.')[-1] == 'docx':
        update.message.reply_text('Print your text')
        return 'docx'
    else:
        update.message.reply_text('Please, print correct name of file')
        return 1


def second_response(update, context):
    global archiv
    try:
        data = update.message.text
        print(2)
        with open(f'{archiv}/{context.user_data["file"]}', 'w', encoding='utf-8') as f:
            f.write(data)
        context.user_data.clear()
        update.message.reply_text("You file is successfully saved!")
        return ConversationHandler.END
    except PermissionError:
        update.message.reply_text("You need to connect a directory")
        return ConversationHandler.END


def third_response(update, context):
    global archiv
    try:
        data = update.message.text
        print(2)
        document = Document()
        document.add_heading((context.user_data["file"][:-5]).capitalize(), 0)
        document.add_paragraph(data)
        document.save(f'{archiv}/{context.user_data["file"]}')
        context.user_data.clear()
        update.message.reply_text("You file is successfully saved!")
        return ConversationHandler.END
    except PermissionError:
        update.message.reply_text("You need to connect a directory")
        return ConversationHandler.END


def stop(update, context):
    update.message.reply_text("Recording has been stopped")
    return ConversationHandler.END


def main():
    updater = Updater('5170844453:AAFtKrDwRUmbQHYP-JHHi2OzQN_BfP2hMzc', use_context=True)
    dp = updater.dispatcher
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler('new_text', new_text)],
        states={
            1: [MessageHandler(Filters.text & ~Filters.command, first_response, pass_user_data=True)],
            'txt': [MessageHandler(Filters.text & ~Filters.command, second_response, pass_user_data=True)],
            'docx': [MessageHandler(Filters.text & ~Filters.command, third_response, pass_user_data=True)]
        },
        fallbacks=[CommandHandler('stop', stop)],
    )

    dp.add_handler(conv_handler)
    dp.add_handler(CommandHandler('start', start))
    dp.add_handler(CommandHandler('stop', stop))
    dp.add_handler(CommandHandler('next_page', next_page))
    dp.add_handler(CommandHandler('get', get))
    dp.add_handler(CommandHandler('new_text', new_text))
    dp.add_handler(CommandHandler('help', help))
    dp.add_handler(CommandHandler('connect', connect))
    dp.add_handler(CommandHandler('disconnect', disconnect))
    updater.start_polling()
    updater.idle()


if __name__ == '__main__':
    main()
